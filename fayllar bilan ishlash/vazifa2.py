from docx import Document
# hujjat = Document()
# hujjat.add_heading("Matinlar To'plami-1")
# hujjat.add_paragraph("jbhvsjkmvdk;lrjewklg trjgvmklnwbiovkmwrkvnfoknwm tob tiobwlk btbirt voirgv vtoignv rinvw irgnv owrivn brtv rbtbireoibreimbttibrmiomebtn itnbrirt bnb")

# hujjat.add_heading("Matinlar To'plami-2")
# hujjat.add_paragraph("jbhvsjkmvdk;lrjewklg trjgvmklnwbiovkmwrkvnfoknwm tob tiobwlk btbirt voirgv vtoignv rinvw irgnv owrivn brtv rbtbireoibreimbttibrmiomebtn itnbrirt bnb")

# hujjat.add_heading("Matinlar To'plami-3")
# hujjat.add_paragraph("jbhvsjkmvdk;lrjewklg trjgvmklnwbiovkmwrkvnfoknwm tob tiobwlk btbirt voirgv vtoignv rinvw irgnv owrivn brtv rbtbireoibreimbttibrmiomebtn itnbrirt bnb")

# hujjat.add_heading("Matinlar To'plami-4")
# hujjat.add_paragraph("jbhvsjkmvdk;lrjewklg trjgvmklnwbiovkmwrkvnfoknwm tob tiobwlk btbirt voirgv vtoignv rinvw irgnv owrivn brtv rbtbireoibreimbttibrmiomebtn itnbrirt bnb")

# hujjat.save("C:/Users/Diyorbek/Documents/GitHub.docx")

class word:
    def __init__(self ):
        self.hujjat = Document() 
    def matn_qushish(self,matn):
        self.hujjat.add_paragraph(matn)
    def Sarlavha(self , matn):
        self.hujjat.add_heading(matn) 
    def add_text(self , text) :
        self.hujjat.add_paragraph(text)
    def  add_heading(self,text):
        self.hujjat.add_heading(text)
    # def frist_paragraf(self):
    #     return self.hujjat.add_paragraphs[0]
    def massiv(self):
     a=[]
     for i in self.hujjat.paragraphs:
        a.append(i.text)

     return  a

     def del_paragraph(self, n):
         if 0 <= n < len(self.hujjat.paragraphs):
             del self.hujjat.paragraphs[n]

     def __repr__(self):
         return self.name

     def __call__(self, n):
         return self.hujjat.paragraphs[n].text
    def saqlash(self):
        self.hujjat.save("C:/Users/Diyorbek/Documents/GitHub/python_darslari/fayllar bilan ishlash/word.docx")
obj1 = word()
obj1.matn_qushish("salom alaykum ")
obj1.Sarlavha("uzbek")
obj1.add_text("Salom")
obj1.add_heading("Diyorbek")
# obj1.frist_paragraf()
obj1.massiv()
# obj1.del_paragraph(1)
obj1
# obj1()
obj1.saqlash()